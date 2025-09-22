#!/usr/bin/env python3
"""
Modul untuk mencetak kuitansi BOS menggunakan python-docx
Template .docx akan digunakan dan dicetak langsung tanpa menyimpan file
"""

import os
import sys
import tempfile
import subprocess
from datetime import datetime
from docx import Document
from docx.api import Document as DocxRaw
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from PySide6.QtWidgets import QMessageBox
from modules.sekolah import getIdentitas
from modules.utils import terbilang

def cetakTransaksiDocx(parent, data):
    """Cetak kuitansi tunggal menggunakan template docx"""
    try:
        # Path template - handle both development dan PyInstaller bundle
        if hasattr(sys, '_MEIPASS'):
            # Running as PyInstaller bundle
            base_path = sys._MEIPASS
            template_path = os.path.join(base_path, "templates", "kuitansi_template.docx")
        else:
            # Running in development mode
            current_dir = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(current_dir, "..", "templates", "kuitansi_template.docx")

        if not os.path.exists(template_path):
            QMessageBox.critical(parent, "Error", f"Template tidak ditemukan: {template_path}")
            return

        # Generate kuitansi
        doc_filled = fill_template_with_data(template_path, data)

        if doc_filled:
            # Print langsung
            success = print_docx_direct(doc_filled, parent)
            if success:
                QMessageBox.information(parent, "Info", "Kuitansi berhasil dicetak!")
            else:
                QMessageBox.critical(parent, "Error", "Gagal mencetak kuitansi!")

    except Exception as e:
        QMessageBox.critical(parent, "Error", f"Gagal mencetak kuitansi: {str(e)}")

def cetakSemuaTransaksiDocx(parent, data_list):
    """Cetak semua kuitansi menggunakan template docx"""
    try:
        if not data_list:
            QMessageBox.warning(parent, "Warning", "Tidak ada data transaksi untuk dicetak!")
            return

        # Path template - handle both development dan PyInstaller bundle
        if hasattr(sys, '_MEIPASS'):
            # Running as PyInstaller bundle
            base_path = sys._MEIPASS
            template_path = os.path.join(base_path, "templates", "kuitansi_template.docx")
        else:
            # Running in development mode
            current_dir = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(current_dir, "..", "templates", "kuitansi_template.docx")

        if not os.path.exists(template_path):
            QMessageBox.critical(parent, "Error", f"Template tidak ditemukan: {template_path}")
            return

        # Generate semua kuitansi dalam satu dokumen
        doc_combined = combine_all_receipts(template_path, data_list)

        if doc_combined:
            # Print langsung
            success = print_docx_direct(doc_combined, parent)
            if success:
                QMessageBox.information(parent, "Info", f"Berhasil mencetak {len(data_list)} kuitansi!")
            else:
                QMessageBox.critical(parent, "Error", "Gagal mencetak kuitansi!")

    except Exception as e:
        QMessageBox.critical(parent, "Error", f"Gagal mencetak semua kuitansi: {str(e)}")

def fill_template_with_data(template_path, data):
    """Fill template docx dengan data transaksi"""
    # try:
    #     # Buka template
    #     doc = Document(template_path)

    #     # Get data sekolah
    #     school_data = getIdentitas()

    #     # Format tanggal
    #     tanggal = data.get('tanggal', '-')
    #     if tanggal != '-':
    #         try:
    #             date_obj = datetime.strptime(tanggal, "%Y-%m-%d")
    #             tanggal_formatted = date_obj.strftime("%d %B %Y")
    #             # Translate month names to Indonesian
    #             bulan_indo = {
    #                 'January': 'Januari', 'February': 'Februari', 'March': 'Maret',
    #                 'April': 'April', 'May': 'Mei', 'June': 'Juni',
    #                 'July': 'Juli', 'August': 'Agustus', 'September': 'September',
    #                 'October': 'Oktober', 'November': 'November', 'December': 'Desember'
    #             }
    #             for eng, indo in bulan_indo.items():
    #                 tanggal_formatted = tanggal_formatted.replace(eng, indo)
    #         except:
    #             tanggal_formatted = tanggal
    #     else:
    #         tanggal_formatted = tanggal

    #     # Format nominal
    #     nominal_raw = data.get('nominal', data.get('nilai', '0'))
    #     if isinstance(nominal_raw, (int, float)):
    #         nominal_clean = str(int(nominal_raw))
    #     else:
    #         nominal_clean = str(nominal_raw).replace('.', '').replace(',', '')

    #     # Format nominal dengan pemisah ribuan
    #     try:
    #         nominal_formatted = '{:,.0f}'.format(int(nominal_clean)).replace(',', '.')
    #     except:
    #         nominal_formatted = nominal_clean

    #     # Data untuk replacement
    #     replacements = {
    #         '{{kabupaten}}': school_data.get('kabupaten', 'KABUPATEN').upper() if school_data else 'KABUPATEN',
    #         '{{kecamatan}}': school_data.get('kecamatan', 'KECAMATAN').upper() if school_data else 'KECAMATAN',
    #         '{{sekolah}}': school_data.get('sekolah', 'NAMA SEKOLAH').upper() if school_data else 'NAMA SEKOLAH',
    #         '{{npsn}}': school_data.get('npsn', '00000000') if school_data else '00000000',
    #         '{{alamat}}': school_data.get('alamat', 'Alamat Sekolah') if school_data else 'Alamat Sekolah',
    #         '{{desa}}': school_data.get('desa', 'Desa') if school_data else 'Desa',
    #         '{{kode_pos}}': school_data.get('kode_pos', '00000') if school_data else '00000',
    #         '{{telp}}': school_data.get('telp', '-') if school_data else '-',
    #         '{{email}}': school_data.get('email', '-') if school_data else '-',
    #         '{{no_bukti}}': data.get('no_bukti', '-'),
    #         '{{tanggal}}': tanggal_formatted,
    #         '{{kode_rekening}}': data.get('kode_rekening', '-'),
    #         '{{kode_kegiatan}}': data.get('kode_kegiatan', '-'),
    #         '{{uraian}}': data.get('uraian', '-'),
    #         '{{terbilang}}': terbilang(nominal_clean).title(),
    #         '{{nominal}}': nominal_formatted,
    #         '{{ks}}': school_data.get('ks', 'NAMA KEPALA SEKOLAH') if school_data else 'NAMA KEPALA SEKOLAH',
    #         '{{bendahara}}': school_data.get('bendahara', 'NAMA BENDAHARA') if school_data else 'NAMA BENDAHARA',
    #         '{{nip_ks}}': school_data.get('nip_ks', '000000000000000000') if school_data else '000000000000000000',
    #         '{{nip_bendahara}}': school_data.get('nip_bendahara', '000000000000000000') if school_data else '000000000000000000',
    #     }

    #     # Replace text dalam paragraphs
    #     for paragraph in doc.paragraphs:
    #         for placeholder, value in replacements.items():
    #             if placeholder in paragraph.text:
    #                 paragraph.text = paragraph.text.replace(placeholder, value)

    #     # Replace text dalam tables
    #     for table in doc.tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 for paragraph in cell.paragraphs:
    #                     for placeholder, value in replacements.items():
    #                         if placeholder in paragraph.text:
    #                             paragraph.text = paragraph.text.replace(placeholder, value)

    #     return doc

    try:
        doc = DocxTemplate(template_path)
        school_data = getIdentitas()
        tanggal = data.get('tanggal', '-')
        if tanggal != '-':
            try:
                date_obj = datetime.strptime(tanggal, '%Y-%m-%d')
                tanggal_formatted = date_obj.strftime('%d %B %Y')
                bulan_indo = {
                    'January': 'Januari', 'February': 'Februari', 'March': 'Maret',
                    'April': 'April', 'May': 'Mei', 'June': 'Juni',
                    'July': 'Juli', 'August': 'Agustus', 'September': 'September',
                    'October': 'Oktober', 'November': 'November', 'December': 'Desember'
                }
                for eng, indo in bulan_indo.items():
                    tanggal_formatted = tanggal_formatted.replace(eng, indo)
            except:
                tanggal_formatted = tanggal
        else:
            tanggal_formatted = tanggal

        #NOminal
        nominal_raw = data.get('nominal', data.get('nilai','0'))
        if isinstance(nominal_raw, (int, float)):
            nominal_clean = str(int(nominal_raw))
        else:
            nominal_clean = str(nominal_raw).replace('.','').replace(',','')

        try:
            nominal_formatted = '{:,.0f}'.format(int(nominal_clean)).replace(',', '.')
        except:
            nominal_formatted = nominal_clean

        #Context data
        context = {
            'tahun': datetime.strptime(data.get('tanggal', ''), '%Y-%m-%d').strftime('%Y'),
            'kabupaten': school_data.get('kabupaten', 'KABUPATEN').upper() if school_data else 'KABUPATEN',
            'kecamatan': school_data.get('kecamatan', 'KECAMATAN').upper() if school_data else 'KECAMATAN',
            'sekolah': school_data.get('sekolah', 'NAMA SEKOLAH').upper() if school_data else 'NAMA SEKOLAH',
            'npsn': school_data.get('npsn', '00000000') if school_data else '00000000',
            'alamat': school_data.get('alamat', 'Alamat Sekolah') if school_data else 'Alamat Sekolah',
            'desa': school_data.get('desa', 'Desa') if school_data else 'Desa',
            'kode_pos': school_data.get('kode_pos', '00000') if school_data else '00000',
            'telp': school_data.get('telp', '-') if school_data else '-',
            'email': school_data.get('email', '-') if school_data else '-',
            'no_bukti': data.get('no_bukti', '-'),
            'tanggal': tanggal_formatted,
            'kode_rekening': data.get('kode_rekening', '-'),
            'kode_kegiatan': data.get('kode_kegiatan', '-'),
            'uraian': data.get('uraian', '-'),
            'terbilang': terbilang(nominal_clean).title(),
            'nominal': nominal_formatted,
            'ks': school_data.get('ks', 'NAMA KEPALA SEKOLAH') if school_data else 'NAMA KEPALA SEKOLAH',
            'bendahara': school_data.get('bendahara', 'NAMA BENDAHARA') if school_data else 'NAMA BENDAHARA',
            'nip_ks': school_data.get('nip_ks', '000000000000000000') if school_data else '000000000000000000',
            'nip_bendahara': school_data.get('nip_bendahara', '000000000000000000') if school_data else '000000000000000000',
        }

        doc.render(context)
        return doc
    except Exception as e:
            print(f"Error filling template: {str(e)}")
            return None

def combine_all_receipts(template_path, data_list):
    """Gabungkan semua kuitansi dalam satu dokumen dengan page break"""
    # try:
    #     combined_doc = None

    #     for index, data in enumerate(data_list):
    #         # Fill template untuk setiap data
    #         doc = fill_template_with_data(template_path, data)

    #         if doc is None:
    #             continue

    #         if combined_doc is None:
    #             # Dokumen pertama
    #             combined_doc = doc
    #         else:
    #             # Tambahkan page break
    #             combined_doc.add_page_break()

    #             # Append semua elemen dari doc ke combined_doc
    #             for element in doc.element.body:
    #                 combined_doc.element.body.append(element)

    #     return combined_doc


    try:
        if not data_list:
            print("No data to combine")
            return None

        first_doc = fill_template_with_data(template_path, data_list[0])
        if first_doc is None:
            return None

        temp_dir = tempfile.mkdtemp()
        combined_path = os.path.join(temp_dir, "combined.docx")
        first_path = os.path.join(temp_dir, "first.docx")
        first_doc.save(first_path)

        master = DocxRaw(first_path)
        composer = Composer(master)

        for data in data_list[1:]:
            doc = fill_template_with_data(template_path, data)
            if doc is None:
                continue
            temp_doc_path = os.path.join(temp_dir, f"temp_{data.get('no_bukti','x')}.docx")
            doc.save(temp_doc_path)
            composer.append(DocxRaw(temp_doc_path))

        composer.save(combined_path)
        return DocxRaw(combined_path)
    except Exception as e:
        print(f"Error combining receipts: {str(e)}")
        return None

def print_docx_direct(doc, parent):
    """Print dokumen Word langsung tanpa menyimpan file permanen"""
    try:
        # Buat temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            temp_path = tmp_file.name

        # Simpan dokumen ke temporary file
        doc.save(temp_path)

        # Coba beberapa aplikasi untuk membuka file .docx di macOS
        apps_to_try = [
            ['open', '-a', 'Microsoft Word', temp_path],
            ['open', '-a', 'Pages', temp_path],
            ['open', '-a', 'LibreOffice', temp_path],
            ['open', temp_path]  # Default app
        ]

        success = False
        app_used = ""

        for app_cmd in apps_to_try:
            try:
                print(f"Trying to open with: {' '.join(app_cmd)}")
                result = subprocess.run(app_cmd, check=True, capture_output=True, text=True, timeout=10)

                # Tunggu sebentar untuk aplikasi terbuka
                import time
                time.sleep(2)

                app_used = app_cmd[2] if len(app_cmd) > 2 else "default application"
                success = True
                break

            except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
                print(f"Failed to open with {app_cmd}: {e}")
                continue

        if success:
            # Beri pesan bahwa user bisa print manual
            from PySide6.QtWidgets import QMessageBox
            reply = QMessageBox.question(parent, "Print",
                                       f"Dokumen telah dibuka di {app_used}.\\n\\n"
                                       "Silahkan tekan Cmd+P untuk mencetak atau tutup aplikasi jika tidak ingin mencetak.\\n\\n"
                                       "Klik OK jika sudah selesai.",
                                       QMessageBox.StandardButton.Ok)
            return True
        else:
            # Jika semua gagal, coba convert ke PDF dan buka di Preview
            return convert_and_print_pdf(doc, parent, temp_path)

    except Exception as e:
        print(f"Error in direct print: {str(e)}")
        return False

    finally:
        # Cleanup temporary file setelah delay
        try:
            import threading
            def cleanup():
                import time
                time.sleep(10)  # Tunggu 10 detik sebelum hapus
                try:
                    os.unlink(temp_path)
                except:
                    pass

            cleanup_thread = threading.Thread(target=cleanup)
            cleanup_thread.daemon = True
            cleanup_thread.start()

        except:
            pass

def print_docx_with_lpr(doc, parent):
    """Alternative: Print menggunakan lpr command (macOS/Linux)"""
    try:
        # Buat temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            temp_path = tmp_file.name

        # Simpan dokumen ke temporary file
        doc.save(temp_path)

        # Convert ke PDF terlebih dahulu (jika ada pandoc atau libreoffice)
        pdf_path = temp_path.replace('.docx', '.pdf')

        try:
            # Coba gunakan LibreOffice untuk convert ke PDF
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path), temp_path
            ], capture_output=True, text=True, timeout=30)

            if result.returncode == 0 and os.path.exists(pdf_path):
                # Print PDF menggunakan lpr
                subprocess.run(['lpr', pdf_path], check=True)
                return True
            else:
                print("LibreOffice convert failed, fallback to Preview method")
                return print_docx_direct(doc, parent)

        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
            print("LibreOffice not found, fallback to Preview method")
            return print_docx_direct(doc, parent)

        finally:
            # Cleanup
            try:
                os.unlink(temp_path)
                if os.path.exists(pdf_path):
                    os.unlink(pdf_path)
            except:
                pass

    except Exception as e:
        print(f"Error in lpr print: {str(e)}")
        return False

def convert_and_print_pdf(doc, parent, temp_docx_path):
    """Convert DOCX to PDF dan buka di Preview sebagai fallback"""
    try:
        print("Trying to convert DOCX to PDF...")

        # Path PDF temporary
        pdf_path = temp_docx_path.replace('.docx', '.pdf')

        # Coba beberapa method untuk convert DOCX ke PDF
        conversion_methods = [
            # Method 1: LibreOffice
            lambda: subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path), temp_docx_path
            ], check=True, capture_output=True, timeout=30),

            # Method 2: pandoc
            lambda: subprocess.run([
                'pandoc', temp_docx_path, '-o', pdf_path
            ], check=True, capture_output=True, timeout=30),
        ]

        conversion_success = False

        for method in conversion_methods:
            try:
                method()
                if os.path.exists(pdf_path):
                    conversion_success = True
                    print(f"Successfully converted to PDF: {pdf_path}")
                    break
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as e:
                print(f"Conversion method failed: {e}")
                continue

        if conversion_success:
            # Buka PDF di Preview
            try:
                subprocess.run(['open', '-a', 'Preview', pdf_path], check=True)

                # Tunggu sebentar untuk aplikasi terbuka
                import time
                time.sleep(2)

                # Beri pesan bahwa user bisa print manual
                from PySide6.QtWidgets import QMessageBox
                reply = QMessageBox.question(parent, "Print",
                                           "Dokumen telah dikonversi ke PDF dan dibuka di Preview.\\n\\n"
                                           "Silahkan tekan Cmd+P untuk mencetak atau tutup Preview jika tidak ingin mencetak.\\n\\n"
                                           "Klik OK jika sudah selesai.",
                                           QMessageBox.StandardButton.Ok)

                # Cleanup PDF file setelah selesai
                import threading
                def cleanup_pdf():
                    time.sleep(10)
                    try:
                        os.unlink(pdf_path)
                    except:
                        pass

                cleanup_thread = threading.Thread(target=cleanup_pdf)
                cleanup_thread.daemon = True
                cleanup_thread.start()

                return True

            except subprocess.CalledProcessError as e:
                print(f"Failed to open PDF in Preview: {e}")
                return False
        else:
            # Jika konversi gagal, berikan pesan error dengan saran
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.warning(parent, "Warning",
                              "Tidak dapat membuka dokumen secara otomatis.\\n\\n"
                              "File telah disimpan di lokasi temporary.\\n"
                              f"Path: {temp_docx_path}\\n\\n"
                              "Silahkan buka file manual dan print dari aplikasi Word/Pages.")
            return False

    except Exception as e:
        print(f"Error in PDF conversion: {str(e)}")
        return False
