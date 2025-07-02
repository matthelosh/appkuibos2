#!/usr/bin/env python3
"""
Script untuk membuat template kuitansi BOS dalam format .docx
Template ini akan digunakan sebagai base untuk generate kuitansi
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_kuitansi_template():
    """Membuat template kuitansi BOS dalam format docx"""
    
    # Buat dokumen baru
    doc = Document()
    
    # Set margin dokumen (kiri=2cm, lainnya=1.5cm)
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.page_height = Cm(33)  # F4/Folio height
        section.page_width = Cm(21)   # F4/Folio width
    
    # HEADER SEKOLAH
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Cm(2)
    header_table.columns[1].width = Cm(15)
    header_table.columns[2].width = Cm(2)
    
    # Header content di kolom tengah
    header_cell = header_table.cell(0, 1)
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Judul header
    run1 = header_para.add_run("PEMERINTAH {{kabupaten}}")
    run1.font.size = Pt(16)
    run1.font.bold = True
    header_para.add_run("\n")
    
    run2 = header_para.add_run("DINAS PENDIDIKAN")
    run2.font.size = Pt(16)
    run2.font.bold = True
    header_para.add_run("\n")
    
    run3 = header_para.add_run("KORWIL KECAMATAN {{kecamatan}}")
    run3.font.size = Pt(16)
    run3.font.bold = True
    header_para.add_run("\n")
    
    run4 = header_para.add_run("{{sekolah}}")
    run4.font.size = Pt(16)
    run4.font.bold = True
    header_para.add_run("\n")
    
    run5 = header_para.add_run("NPSN: {{npsn}}")
    run5.font.size = Pt(14)
    header_para.add_run("\n")
    
    run6 = header_para.add_run("{{alamat}}, {{desa}}, {{kecamatan}}, Kode Pos {{kode_pos}}")
    run6.font.size = Pt(12)
    header_para.add_run("\n")
    
    run7 = header_para.add_run("Telp.: {{telp}}, Email: {{email}}")
    run7.font.size = Pt(12)
    
    # Tambah border bawah pada header
    add_border_bottom(header_table)
    
    # Spasi
    doc.add_paragraph()
    
    # TABEL INFO ANGGARAN
    info_table = doc.add_table(rows=3, cols=5)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Row 1
    info_table.cell(0, 0).text = "Mata Anggaran"
    info_table.cell(0, 1).text = ": BOS Reguler 2024"
    info_table.cell(0, 3).text = "No Bukti"
    info_table.cell(0, 4).text = ": {{no_bukti}}"
    
    # Row 2
    info_table.cell(1, 0).text = "Tahun Anggaran"
    info_table.cell(1, 1).text = ": Tahun 2024"
    info_table.cell(1, 3).text = "Tanggal"
    info_table.cell(1, 4).text = ": {{tanggal}}"
    
    # Row 3
    info_table.cell(2, 0).text = "Sumber Dana"
    info_table.cell(2, 1).text = ": APBN"
    
    # Set lebar kolom
    info_table.columns[0].width = Cm(3)
    info_table.columns[1].width = Cm(4)
    info_table.columns[2].width = Cm(3)  # Spacer
    info_table.columns[3].width = Cm(3)
    info_table.columns[4].width = Cm(4)
    
    # Spasi
    doc.add_paragraph()
    
    # JUDUL KUITANSI
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("KUITANSI")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.underline = True
    
    # Spasi
    doc.add_paragraph()
    
    # KODE REKENING DAN KEGIATAN
    code_para = doc.add_paragraph()
    code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_run1 = code_para.add_run("Kode Rekening: {{kode_rekening}}")
    code_run1.font.bold = True
    code_para.add_run("          ")
    code_run2 = code_para.add_run("Kode Kegiatan: {{kode_kegiatan}}")
    code_run2.font.bold = True
    
    # Spasi
    doc.add_paragraph()
    
    # TABEL KONTEN UTAMA
    content_table = doc.add_table(rows=3, cols=2)
    content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    content_table.style = 'Table Grid'
    
    # Set lebar kolom
    content_table.columns[0].width = Cm(5)
    content_table.columns[1].width = Cm(12)
    
    # Row 1
    content_table.cell(0, 0).text = "Sudah diterima dari"
    content_table.cell(0, 1).text = "Bendahara BOS Reguler {{sekolah}}"
    
    # Row 2
    content_table.cell(1, 0).text = "Uang Sebesar"
    cell_1_1 = content_table.cell(1, 1)
    para_1_1 = cell_1_1.paragraphs[0]
    run_1_1 = para_1_1.add_run("{{terbilang}} Rupiah")
    run_1_1.italic = True
    
    # Row 3
    content_table.cell(2, 0).text = "Untuk Keperluan"
    content_table.cell(2, 1).text = "{{uraian}}"
    
    # Spasi
    doc.add_paragraph()
    
    # BERKAS TERLAMPIR
    berkas_para = doc.add_paragraph("*) Bukti / Berkas Terlampir")
    berkas_run = berkas_para.runs[0]
    berkas_run.font.bold = True
    
    # Spasi
    doc.add_paragraph()
    
    # NOMINAL BESAR
    nominal_para = doc.add_paragraph()
    nominal_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nominal_run = nominal_para.add_run("Rp. {{nominal}},-")
    nominal_run.font.size = Pt(18)
    nominal_run.font.bold = True
    
    # Spasi besar
    doc.add_paragraph()
    doc.add_paragraph()
    
    # TANGGAL DAN TEMPAT
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run("{{kabupaten}}, {{tanggal}}")
    
    # Spasi
    doc.add_paragraph()
    
    # TABEL TANDA TANGAN
    ttd_table = doc.add_table(rows=4, cols=3)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set lebar kolom
    for col in ttd_table.columns:
        col.width = Cm(6)
    
    # Row 1: Jabatan
    ttd_table.cell(0, 0).text = "Menyetujui,"
    ttd_table.cell(0, 1).text = "Yang mengeluarkan,"
    ttd_table.cell(0, 2).text = "Yang menerima,"
    
    # Row 2: Sub jabatan
    ttd_table.cell(1, 0).text = "Kepala Sekolah"
    ttd_table.cell(1, 1).text = "Bendahara"
    ttd_table.cell(1, 2).text = "Pelaksana"
    
    # Row 3: Nama (dengan ruang untuk tanda tangan)
    ttd_table.cell(2, 0).text = "\n\n\n{{ks}}"
    ttd_table.cell(2, 1).text = "\n\n\n{{bendahara}}"
    ttd_table.cell(2, 2).text = "\n\n\n..............................."
    
    # Row 4: NIP
    ttd_table.cell(3, 0).text = "NIP. {{nip_ks}}"
    ttd_table.cell(3, 1).text = "NIP. {{nip_bendahara}}"
    ttd_table.cell(3, 2).text = "NIP. "
    
    # Center align semua cell tanda tangan
    for row in ttd_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Buat nama tebal dan garis bawah pada row 3
    for i in range(2):
        cell = ttd_table.cell(2, i)
        para = cell.paragraphs[0]
        para.clear()
        para.add_run("\n\n\n")
        name_run = para.add_run(f"{{{{'ks' if i == 0 else 'bendahara'}}}}")
        name_run.font.bold = True
        name_run.underline = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def add_border_bottom(table):
    """Menambahkan border bawah pada tabel"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    
    tblBorders.append(bottom)
    tblPr.append(tblBorders)

if __name__ == "__main__":
    # Buat template dan simpan
    template = create_kuitansi_template()
    template.save("/Users/matthelosh/Projects/piton/appkuibos2/templates/kuitansi_template.docx")
    print("Template kuitansi berhasil dibuat: templates/kuitansi_template.docx")
